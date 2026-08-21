"""Proves the layout-aware-ingestion-and-vision milestone's DOCX and PDF paths end to end.

Self-contained by design (matching test_structured_markdown_ingestion.py's
precedent): builds a small synthetic .docx/.pdf at test time. never the
real, git-ignored data/knowledge_base. ingests it into a fresh
pytest-namespaced dataset_id, retrieves against it through a real
Postgres round trip, and cleans up afterward.

The PDF fixtures are built with `reportlab` (a `dev`-extras-only test
dependency. pdfplumber/pypdf are both PDF *readers*; this is the one PDF
*writer* the test suite needs, and it is never imported outside this
file). Its default `getSampleStyleSheet()` gives `Title` and `Heading1`
the same 18pt size, which would collapse both into `PDFLoader`'s single
title tier, so this file defines its own three explicit-size paragraph
styles (`_TITLE_STYLE`/`_HEADING_STYLE`/`_BODY_STYLE`) to keep the two
heading tiers distinct.
"""

from __future__ import annotations

import io
import uuid
from pathlib import Path

import docx
from PIL import Image
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    Image as RLImage,
)
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from rag.api.routers.query import SourceItem
from rag.config import load_config
from rag.ingestion.pipeline import IngestionPipeline
from rag.retrieval.pipeline import RetrievalPipeline, source_dict

_TITLE_STYLE = ParagraphStyle(name="TitleX", fontName="Helvetica-Bold", fontSize=20, leading=24)
_HEADING_STYLE = ParagraphStyle(name="HeadingX", fontName="Helvetica-Bold", fontSize=14, leading=18)
_BODY_STYLE = ParagraphStyle(name="BodyX", fontName="Helvetica", fontSize=10, leading=13)


def _tiny_png_bytes() -> bytes:
    """Build a minimal valid 1x1 PNG, for a real embeddable image."""
    buf = io.BytesIO()
    Image.new("RGB", (1, 1), color="blue").save(buf, format="PNG")
    return buf.getvalue()


def _build_docx(tmp_path: Path) -> Path:
    """Build a synthetic .docx with a heading, a table, a page break, and an embedded image."""
    doc_dir = tmp_path / "layout_vision_extension" / "runbooks"
    doc_dir.mkdir(parents=True)
    (doc_dir / "assets").mkdir()
    (doc_dir / "assets" / "recovery-view-01.png").write_bytes(_tiny_png_bytes())

    path = doc_dir / "recovery-guide.docx"
    document = docx.Document()
    document.add_paragraph("Recovery Guide", style="Title")
    document.add_paragraph("1. Trigger", style="Heading 1")
    document.add_paragraph("Trigger when consumer lag exceeds ten minutes.")
    document.add_page_break()
    document.add_paragraph("2. Diagnosis", style="Heading 1")
    document.add_picture(io.BytesIO(_tiny_png_bytes()))
    document.add_paragraph("Figure 1. Diagnosis flow.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Partition"
    table.cell(0, 1).text = "Lag"
    table.cell(1, 0).text = "p-3"
    table.cell(1, 1).text = "9000"
    document.save(str(path))
    return path


def _ingest(config, tmp_path: Path) -> tuple[IngestionPipeline, str, str]:
    """Ingest the synthetic docx into a fresh dataset_id.

    Returns
    -------
    tuple[IngestionPipeline, str, str]
        ``(pipeline, dataset_id, document_id)``.
    """
    dataset_id = f"pytest-layout-vision-{uuid.uuid4()}"
    doc_path = _build_docx(tmp_path)
    pipeline = IngestionPipeline(config)
    result = pipeline.ingest_file(doc_path, dataset_id)
    return pipeline, dataset_id, result["document_id"]


def test_docx_table_row_retrieves_table_chunk_with_page_and_section(
    require_postgres, config, tmp_path: Path
):
    """A table-row question retrieves the table chunk, correctly tagged with page/section_path."""
    pipeline, dataset_id, document_id = _ingest(config, tmp_path)
    retrieval = RetrievalPipeline(config)
    try:
        results = retrieval.retrieve(
            "What is the lag for partition p-3?",
            filters={"dataset_id": dataset_id},
            candidate_k=10,
            generation_context_top_n=10,
        )
        table_results = [r for r in results if r.chunk.metadata.content_type == "table"]
        assert table_results, "expected a table chunk in top-10 retrieval"
        table_result = table_results[0]
        assert "p-3" in table_result.chunk.content
        assert "9000" in table_result.chunk.content
        # The table sits after the page break, in section "2. Diagnosis".
        assert table_result.chunk.metadata.page == 2
        assert table_result.chunk.metadata.section_path is not None
        assert "2. Diagnosis" in table_result.chunk.metadata.section_path
    finally:
        pipeline._vectorstore.delete_document(document_id)


def test_docx_image_resolves_to_existing_asset_with_correct_page(
    require_postgres, config, tmp_path: Path
):
    """The embedded image's source_anchor resolves to the pre-existing asset, with its page."""
    pipeline, dataset_id, document_id = _ingest(config, tmp_path)
    retrieval = RetrievalPipeline(config)
    try:
        results = retrieval.retrieve(
            "diagnosis flow diagram",
            filters={"dataset_id": dataset_id},
            candidate_k=10,
            generation_context_top_n=10,
        )
        image_results = [r for r in results if r.chunk.metadata.content_type == "image"]
        assert image_results, "expected an image chunk in top-10 retrieval"
        image_result = image_results[0]
        assert image_result.chunk.metadata.source_anchor == "assets/recovery-view-01.png"
        assert image_result.chunk.metadata.page == 2
        assert "Figure 1. Diagnosis flow." in image_result.chunk.content
    finally:
        pipeline._vectorstore.delete_document(document_id)


def test_docx_heading_before_page_break_reports_page_one(require_postgres, config, tmp_path: Path):
    """Content before the manual page break is tagged page=1, not 2."""
    pipeline, dataset_id, document_id = _ingest(config, tmp_path)
    retrieval = RetrievalPipeline(config)
    try:
        results = retrieval.retrieve(
            "When should an operator trigger this guide?",
            filters={"dataset_id": dataset_id},
            candidate_k=10,
            generation_context_top_n=10,
        )
        trigger_results = [r for r in results if "consumer lag" in r.chunk.content]
        assert trigger_results, "expected the trigger paragraph in top-10 retrieval"
        assert trigger_results[0].chunk.metadata.page == 1
    finally:
        pipeline._vectorstore.delete_document(document_id)


def test_incremental_reingestion_of_unchanged_docx_is_a_noop(
    require_postgres, config, tmp_path: Path
):
    """Re-ingesting the same unchanged .docx skips re-writing chunks (checksum-gated)."""
    pipeline, dataset_id, document_id = _ingest(config, tmp_path)
    try:
        doc_path = tmp_path / "layout_vision_extension" / "runbooks" / "recovery-guide.docx"
        result = pipeline.ingest_file(doc_path, dataset_id)
        assert result["changed"] is False
        assert result["document_id"] == document_id
    finally:
        pipeline._vectorstore.delete_document(document_id)


def _build_pdf(tmp_path: Path) -> Path:
    """Build a synthetic, multi-page .pdf mirroring `_build_docx`'s structure.

    title -> heading -> prose -> page break -> heading -> embedded
    image+caption -> table, so the same normalization contract (page/
    section_path/content_type/parent_chunk_id) can be asserted on both
    formats together in one test.
    """
    doc_dir = tmp_path / "layout_vision_extension" / "operations"
    doc_dir.mkdir(parents=True)
    image_path = doc_dir / "diagnosis-flow.png"
    image_path.write_bytes(_tiny_png_bytes())

    path = doc_dir / "recovery-guide.pdf"
    story = [
        Paragraph("Recovery Guide", _TITLE_STYLE),
        Spacer(1, 12),
        Paragraph("1. Trigger", _HEADING_STYLE),
        Paragraph("Trigger when consumer lag exceeds ten minutes.", _BODY_STYLE),
        PageBreak(),
        Paragraph("2. Diagnosis", _HEADING_STYLE),
        RLImage(str(image_path), width=100, height=100),
        Paragraph("Figure 1. Diagnosis flow.", _BODY_STYLE),
        Spacer(1, 12),
        Table(
            [["Partition", "Lag"], ["p-3", "9000"]],
            style=TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 1, (0, 0, 0)),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                ]
            ),
        ),
    ]
    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)
    return path


def _ingest_pdf(config, tmp_path: Path) -> tuple[IngestionPipeline, str, str]:
    """Ingest the synthetic PDF into a fresh dataset_id, same contract as `_ingest`."""
    dataset_id = f"pytest-layout-vision-pdf-{uuid.uuid4()}"
    doc_path = _build_pdf(tmp_path)
    pipeline = IngestionPipeline(config)
    result = pipeline.ingest_file(doc_path, dataset_id)
    return pipeline, dataset_id, result["document_id"]


def test_pdf_multi_page_document_preserves_structural_fields_together(
    require_postgres, config, tmp_path: Path
):
    """One multi-page PDF's page/section_path/content_type/parent_chunk_id, asserted together.

    Exercises the whole normalization contract on a single coherent
    document (title -> heading -> prose -> page break -> heading -> image
    -> table) rather than each field in isolation across separate small
    fixtures. the PDF analog of the DOCX table/image/page tests above,
    closing the one format this file previously had no integration
    coverage for (`PDFLoader.load()` end to end, not just its private
    helpers. see `tests/unit/test_loaders.py` for those).
    """
    pipeline, dataset_id, document_id = _ingest_pdf(config, tmp_path)
    retrieval = RetrievalPipeline(config)
    try:
        results = retrieval.retrieve(
            "diagnosis flow diagram partition lag",
            filters={"dataset_id": dataset_id},
            candidate_k=10,
            generation_context_top_n=10,
        )
        by_content_type = {r.chunk.metadata.content_type: r for r in results}

        assert "table" in by_content_type, "expected a table chunk in top-10 retrieval"
        table_result = by_content_type["table"]
        assert "p-3" in table_result.chunk.content
        assert "9000" in table_result.chunk.content
        assert table_result.chunk.metadata.page == 2
        assert table_result.chunk.metadata.section_path is not None
        assert "2. Diagnosis" in table_result.chunk.metadata.section_path

        assert "image" in by_content_type, "expected an image chunk in top-10 retrieval"
        image_result = by_content_type["image"]
        assert image_result.chunk.metadata.page == 2
        assert "Figure 1. Diagnosis flow." in image_result.chunk.content
        assert image_result.chunk.metadata.source_anchor is not None
        assert image_result.chunk.metadata.section_path is not None
        assert "2. Diagnosis" in image_result.chunk.metadata.section_path

        # The "## 2. Diagnosis" heading line has no other text on its own
        # paragraph, so it becomes its own tiny prose span (untagged
        # content_type defaults to "prose"). exactly the "explanatory
        # paragraph introducing this section" _compute_parent_chunk_ids'
        # docstring describes. Both the image and the table that follow it
        # in the same section_path link back to that heading-only chunk,
        # not to each other and not to the unrelated page-1 prose.
        section_parent_id = image_result.chunk.metadata.parent_chunk_id
        assert section_parent_id is not None
        assert table_result.chunk.metadata.parent_chunk_id == section_parent_id

        trigger_results = [r for r in results if "consumer lag" in r.chunk.content]
        assert trigger_results, "expected the page-1 trigger paragraph in top-10 retrieval"
        assert trigger_results[0].chunk.metadata.page == 1
        assert trigger_results[0].chunk.metadata.content_type == "prose"
    finally:
        pipeline._vectorstore.delete_document(document_id)


def test_real_ollama_vision_provider_produces_non_empty_description(
    require_postgres, require_ollama, tmp_path: Path
):
    """A real image sent through a real Ollama/moondream call gets a real, non-empty description.

    Every other vision test in this suite (`test_ollama_vision.py`,
    `test_writer.py`) mocks `ollama.Client`/`VisionProvider`. none of
    them prove `OllamaVisionProvider` actually talks to a real Ollama
    server and gets back something usable. Uses
    `config/experiments/layout-vision-c-vision.yaml` (`vision.provider:
    ollama`, model `moondream`). the exact config the controlled A/B/C
    layout-vision evaluation ran C with. rather than mutating the
    shared `config` fixture, since every other test in this file relies
    on that fixture staying `vision.provider: none`.
    """
    vision_config = load_config("config/experiments/layout-vision-c-vision.yaml")
    doc_dir = tmp_path / "layout_vision_extension" / "architecture"
    (doc_dir / "assets").mkdir(parents=True)
    (doc_dir / "assets" / "diagram.png").write_bytes(_tiny_png_bytes())
    doc_path = doc_dir / "overview.md"
    doc_path.write_text(
        "# Architecture Overview\n\n"
        "![System diagram](assets/diagram.png)\n\n"
        "*Figure 1: System diagram.*\n",
        encoding="utf-8",
    )

    dataset_id = f"pytest-real-vision-{uuid.uuid4()}"
    pipeline = IngestionPipeline(vision_config)
    result = pipeline.ingest_file(doc_path, dataset_id)
    document_id = result["document_id"]
    try:
        retrieval = RetrievalPipeline(vision_config)
        results = retrieval.retrieve(
            "system diagram",
            filters={"dataset_id": dataset_id},
            candidate_k=10,
            generation_context_top_n=10,
        )
        vision_results = [r for r in results if r.chunk.metadata.vision_generated]
        assert vision_results, "expected a vision_generated=True sibling chunk in top-10 retrieval"
        description = vision_results[0].chunk.metadata.vision_description
        assert description is not None
        assert description.strip() != ""
        assert vision_results[0].chunk.content == description
    finally:
        pipeline._vectorstore.delete_document(document_id)


def test_page_survives_insert_through_source_dict_and_source_item(
    require_postgres, config, tmp_path: Path
):
    """`page` survives INSERT -> SELECT/search -> `source_dict` -> `SourceItem` intact.

    Regression test for the real bug this milestone found and fixed
    (`PgVectorStore.add_chunks`'s INSERT column list and `_row_to_metadata`
    never had `page` added, so it silently persisted as NULL for every
    chunk. see ISSUES.md/PROJECT_JOURNAL.md's 2026-08-21 entries). Prior
    coverage stopped at `VectorStore.search()`. this goes one step
    further, through `source_dict` (the exact function `POST /query`'s
    `pipeline.answer()` uses to build its `"sources"` list) and into a
    real `SourceItem` Pydantic model, the actual API response shape a
    caller receives, not just a generic structured-field round trip.
    """
    pipeline, dataset_id, document_id = _ingest(config, tmp_path)
    retrieval = RetrievalPipeline(config)
    try:
        results = retrieval.retrieve(
            "What is the lag for partition p-3?",
            filters={"dataset_id": dataset_id},
            candidate_k=10,
            generation_context_top_n=10,
        )
        table_results = [r for r in results if r.chunk.metadata.content_type == "table"]
        assert table_results, "expected a table chunk in top-10 retrieval"

        rendered = source_dict(table_results[0])
        assert rendered["page"] == 2

        source_item = SourceItem(**rendered)
        assert source_item.page == 2
    finally:
        pipeline._vectorstore.delete_document(document_id)
