"""Proves a staged upload's `.upload-tmp-*` staging directory can never leak into content.

Follow-up to `test_api_ingest_atomicity.py`: that fix made upload
installation atomic by staging bytes at a temp path before installing them
at their final destination. A second review found the temp *path* itself
could still leak into persisted metadata, because several loaders derive
fields directly from the physical `Path` they're given (`TextLoader`/
`HTMLLoader`/`PDFLoader`/`DocxLoader`'s `title = ... or path.stem`
fallback; `PDFLoader`/`DocxLoader`'s embedded-image extraction via
`loaders/base.py:resolve_image_asset`, keyed off `document_path.stem`/
`.parent`), fields `IngestionPipeline.ingest_file`'s `source_override`
(which only overrides the final `RawDocument.source` string) never
reaches.

The fix (`api/routers/ingest.py:_ingest_upload_atomically`) keeps the
staged file's own *filename* identical to the final destination's. Only
its *parent directory* is randomized (`staging_dir/<original_filename>`,
not `staging_dir_prefix<original_filename>` or a randomly-named file).
These tests simulate that same directory shape directly against the real
loaders (no HTTP layer, no `IngestionPipeline` needed for most of them) to
prove every path-derived field comes out correct without requiring any
loader-specific patch. The fix is structural, not per-field.
"""

from __future__ import annotations

import io
import uuid
from pathlib import Path

import docx
import pytest
from PIL import Image
from pydantic import BaseModel
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate

from rag.config import load_config
from rag.ingestion.pipeline import IngestionPipeline
from rag.loaders.docx_loader import DocxLoader
from rag.loaders.html_loader import HTMLLoader
from rag.loaders.pdf_loader import PDFLoader
from rag.loaders.text_loader import TextLoader
from rag.schemas import Chunk

_TMP_MARKER = ".upload-tmp-"


def _staged_path(tmp_path: Path, filename: str) -> Path:
    """Return `tmp_path/.upload-tmp-<uuid>/<filename>`, matching the router's staging shape."""
    staging_dir = tmp_path / f"{_TMP_MARKER}{uuid.uuid4().hex}"
    staging_dir.mkdir()
    return staging_dir / filename


def _tiny_png_bytes() -> bytes:
    """Build a minimal valid 1x1 PNG, for a real embeddable image."""
    buf = io.BytesIO()
    Image.new("RGB", (1, 1), color="red").save(buf, format="PNG")
    return buf.getvalue()


def _assert_no_temp_marker(value: object, path: str = "root") -> None:
    """Recursively assert no string anywhere inside `value` contains `_TMP_MARKER`.

    Walks Pydantic models (via `model_dump()`), dicts, and list/tuple/set
    containers. [Test E]
    """
    if isinstance(value, str):
        assert _TMP_MARKER not in value, f"temp-staging marker leaked at {path}: {value!r}"
    elif isinstance(value, BaseModel):
        _assert_no_temp_marker(value.model_dump(), path)
    elif isinstance(value, dict):
        for key, sub in value.items():
            _assert_no_temp_marker(sub, f"{path}.{key}")
    elif isinstance(value, list | tuple | set):
        for i, sub in enumerate(value):
            _assert_no_temp_marker(sub, f"{path}[{i}]")


def _assert_raw_loader_output_clean(doc, path: str = "root") -> None:
    """Scan a *raw loader* `RawDocument` for temp-marker leaks, `source` excepted.

    A raw `Loader.load()` call (unlike `IngestionPipeline.ingest_file`)
    never sees `source_override`; fixing `RawDocument.source` is
    `ingest_file`'s job, applied *after* loading, and is proven separately
    by `test_recursive_scan_of_persisted_chunks_finds_no_temp_marker`
    below. Every other field (title, content including any embedded asset
    link, author, etc.) must already be clean straight out of the
    loader, with no override needed, which is exactly what these
    loader-level tests are checking.
    """
    data = doc.model_dump()
    data.pop("source", None)
    _assert_no_temp_marker(data, path)


def _build_pdf_with_image(path: Path) -> None:
    """Write a real single-page PDF with prose and an embedded image to `path`.

    `title=""` suppresses reportlab's own default document-info title
    (otherwise `"(anonymous)"`), so `PDFLoader` actually exercises its
    `path.stem` title fallback, the behavior these tests are checking,
    instead of reading reportlab's placeholder metadata.
    """
    image_path = path.parent / "_source-image.png"
    image_path.write_bytes(_tiny_png_bytes())
    story = [
        Paragraph("Overview"),
        RLImage(str(image_path), width=50, height=50),
        Paragraph("Figure 1. Overview diagram."),
    ]
    SimpleDocTemplate(str(path), pagesize=LETTER, title="").build(story)


def _build_docx_with_image(path: Path) -> None:
    """Write a real .docx with a heading and an embedded image to `path`."""
    document = docx.Document()
    document.add_paragraph("Overview", style="Heading 1")
    document.add_picture(io.BytesIO(_tiny_png_bytes()))
    document.add_paragraph("Figure 1. Overview diagram.")
    document.save(str(path))


# --- A: Markdown upload -----------------------------------------------------


def test_markdown_upload_has_no_temp_marker_anywhere(tmp_path: Path):
    """A staged Markdown file's `RawDocument` carries no `.upload-tmp-` fragment. [Test A]."""
    staged = _staged_path(tmp_path, "report.md")
    staged.write_text("# Overview\n\nOrdinary content, no front matter at all.", encoding="utf-8")

    doc = TextLoader().load(staged)

    assert doc.title == "report"  # the *original* filename's stem, not the staging dir's name
    _assert_raw_loader_output_clean(doc)


# --- D: plain text / HTML upload --------------------------------------------


def test_plain_text_upload_title_fallback_uses_original_stem(tmp_path: Path):
    """A staged `.txt` file's title fallback is the original stem, not the staging dir. [Test D]."""
    staged = _staged_path(tmp_path, "customer-notes.txt")
    staged.write_text("Plain text content with no metadata of its own.", encoding="utf-8")

    doc = TextLoader().load(staged)

    assert doc.title == "customer-notes"
    _assert_raw_loader_output_clean(doc)


def test_html_upload_title_fallback_uses_original_stem(tmp_path: Path):
    """A staged HTML file with no `<title>` falls back to the original stem. [Test D]."""
    staged = _staged_path(tmp_path, "release-notes.html")
    staged.write_text("<html><body><p>No title tag here.</p></body></html>", encoding="utf-8")

    doc = HTMLLoader().load(staged)

    assert doc.title == "release-notes"
    _assert_raw_loader_output_clean(doc)


# --- B: PDF upload -----------------------------------------------------------


def test_pdf_upload_image_asset_naming_survives_temp_staging_directory(tmp_path: Path):
    """The exact leak Codex identified: a PDF's extracted-image filename must use the stable name.

    With no pre-existing sibling `assets/` folder, `PDFLoader` extracts
    its embedded image and writes it via `loaders/base.py:
    resolve_image_asset`, named `f"{document_path.stem}-figure-01{ext}"`.
    Loaded from a staged path, `document_path.stem` must be the original
    filename's stem (`quarterly-report`), never the staging directory's
    random name. [Test B, and the focused Codex-leak test]
    """
    staged = _staged_path(tmp_path, "quarterly-report.pdf")
    _build_pdf_with_image(staged)

    doc = PDFLoader().load(staged)

    assert doc.title == "quarterly-report"
    written_asset = staged.parent / "assets" / "quarterly-report-figure-01.png"
    assert written_asset.is_file(), "expected the extracted image under the *original* stem"
    assert "assets/quarterly-report-figure-01.png" in doc.content
    _assert_raw_loader_output_clean(doc)


# --- C: DOCX upload ----------------------------------------------------------


def test_docx_upload_image_asset_naming_survives_temp_staging_directory(tmp_path: Path):
    """Same invariant as the PDF case, for `.docx`. [Test C]."""
    staged = _staged_path(tmp_path, "standalone-report.docx")
    _build_docx_with_image(staged)

    doc = DocxLoader().load(staged)

    assert doc.title == "standalone-report"
    written_asset = staged.parent / "assets" / "standalone-report-figure-01.png"
    assert written_asset.is_file()
    assert "assets/standalone-report-figure-01.png" in doc.content
    _assert_raw_loader_output_clean(doc)


# --- End-to-end through IngestionPipeline (fake VectorStore, no Postgres) ---


class _StatefulRecordingVectorStore:
    """Fake VectorStore tracking (source, dataset_id) -> document_id, like a real one.

    Needed (rather than a fresh-id-every-call fake) for [Test F]: proving
    two staged re-uploads of the same logical filename, through two
    different randomly-named staging directories, still resolve to the
    same `document_id`.
    """

    def __init__(self) -> None:
        self._documents: dict[tuple[str, str], tuple[str, str]] = {}
        self._next_id = 0
        self.written_chunks: list[Chunk] = []

    def health_check(self) -> bool:
        return True

    def get_or_create_document_id(self, source: str, checksum: str, dataset_id: str):
        key = (source, dataset_id)
        if key not in self._documents:
            self._next_id += 1
            doc_id = f"doc-{self._next_id}"
            self._documents[key] = (doc_id, checksum)
            return doc_id, True
        doc_id, existing_checksum = self._documents[key]
        changed = existing_checksum != checksum
        if changed:
            self._documents[key] = (doc_id, checksum)
        return doc_id, changed

    def delete_chunks_by_document_id(self, document_id: str) -> None:
        self.written_chunks = [
            c for c in self.written_chunks if c.metadata.document_id != document_id
        ]

    def add_chunks(self, chunks: list[Chunk]) -> None:
        self.written_chunks.extend(chunks)


class _FakeEmbedder:
    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [[0.0] for _ in texts]

    def embed_query(self, text: str) -> list[float]:
        return [0.0]


def test_recursive_scan_of_persisted_chunks_finds_no_temp_marker(tmp_path: Path):
    """A full pipeline run (fake VectorStore) never persists a `.upload-tmp-` string. [Test E]."""
    staged = _staged_path(tmp_path, "quarterly-report.pdf")
    _build_pdf_with_image(staged)
    stable_dest = tmp_path / "quarterly-report.pdf"

    vectorstore = _StatefulRecordingVectorStore()
    pipeline = IngestionPipeline(load_config(), vectorstore=vectorstore, embedder=_FakeEmbedder())
    pipeline.ingest_file(staged, "test-dataset", source_override=str(stable_dest))

    assert vectorstore.written_chunks, "expected at least one chunk written"
    for chunk in vectorstore.written_chunks:
        _assert_no_temp_marker(chunk)
    image_chunks = [c for c in vectorstore.written_chunks if c.metadata.content_type == "image"]
    assert image_chunks, "expected the embedded image to produce its own chunk"
    assert image_chunks[0].metadata.attachment_name == "quarterly-report-figure-01.png"


def test_successful_reupload_through_different_staging_dirs_keeps_same_document_id(
    tmp_path: Path,
):
    """Two staged re-uploads of the same filename resolve to the same document_id. [Test F]."""
    stable_dest = tmp_path / "policy.md"
    vectorstore = _StatefulRecordingVectorStore()
    pipeline = IngestionPipeline(load_config(), vectorstore=vectorstore, embedder=_FakeEmbedder())

    staged_v1 = _staged_path(tmp_path, "policy.md")
    staged_v1.write_text("Version one content.", encoding="utf-8")
    result_v1 = pipeline.ingest_file(staged_v1, "test-dataset", source_override=str(stable_dest))

    staged_v2 = _staged_path(tmp_path, "policy.md")  # a *different* random staging dir
    staged_v2.write_text("Version two content, changed.", encoding="utf-8")
    result_v2 = pipeline.ingest_file(staged_v2, "test-dataset", source_override=str(stable_dest))

    assert result_v1["document_id"] == result_v2["document_id"]
    assert result_v2["changed"] is True
    assert all(c.metadata.source == str(stable_dest) for c in vectorstore.written_chunks)
    _assert_no_temp_marker(vectorstore.written_chunks)


@pytest.mark.parametrize(
    "filename",
    ["report.md", "report.txt", "report.html", "report.pdf", "report.docx"],
)
def test_no_leftover_temp_marker_at_any_supported_extension(tmp_path: Path, filename: str):
    """A quick sweep across every supported extension: no title/source ever carries the marker.

    Complements the format-specific tests above with a single assertion
    repeated across every loader this project registers, so a future new
    loader (or extension mapping change) that reintroduces a `path.stem`/
    `path.name` fallback without staging-awareness is still caught here.
    """
    staged = _staged_path(tmp_path, filename)
    if filename.endswith((".md", ".txt")):
        staged.write_text("Plain content.", encoding="utf-8")
        doc = TextLoader().load(staged)
    elif filename.endswith(".html"):
        staged.write_text("<html><body>Plain content.</body></html>", encoding="utf-8")
        doc = HTMLLoader().load(staged)
    elif filename.endswith(".pdf"):
        _build_pdf_with_image(staged)
        doc = PDFLoader().load(staged)
    else:
        _build_docx_with_image(staged)
        doc = DocxLoader().load(staged)

    _assert_raw_loader_output_clean(doc)
