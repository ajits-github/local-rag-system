from __future__ import annotations

import io
from pathlib import Path

import docx
import pytest
from PIL import Image
from pypdf import PdfWriter

from rag.chunkers.structured_markdown import StructuredMarkdownChunker
from rag.loaders.base import resolve_image_asset
from rag.loaders.docx_loader import DocxLoader
from rag.loaders.html_loader import HTMLLoader
from rag.loaders.pdf_loader import PDFLoader
from rag.loaders.registry import get_loader
from rag.loaders.text_loader import TextLoader


def _tiny_png_bytes() -> bytes:
    """Build a minimal valid 1x1 PNG, for tests that need a real embeddable image."""
    buf = io.BytesIO()
    Image.new("RGB", (1, 1), color="red").save(buf, format="PNG")
    return buf.getvalue()


def test_text_loader_reads_plain_text(tmp_path: Path):
    """A plain .txt file is read as-is with source_type "text"."""
    path = tmp_path / "note.txt"
    path.write_text("hello world", encoding="utf-8")

    doc = TextLoader().load(path)

    assert doc.content == "hello world"
    assert doc.source_type == "text"
    assert doc.source == str(path)


def test_text_loader_parses_yaml_frontmatter(tmp_path: Path):
    """YAML front-matter populates title/author/last_modified and is stripped from content."""
    path = tmp_path / "policy.md"
    path.write_text(
        "---\n"
        'title: "Access Control Policy"\n'
        'owner: "Security Engineering"\n'
        'last_reviewed: "2026-07-15"\n'
        'tags: ["access", "rbac"]\n'
        "---\n\n"
        "# Access Control Policy\n\nBody text here.",
        encoding="utf-8",
    )

    doc = TextLoader().load(path)

    assert doc.title == "Access Control Policy"
    assert doc.author == "Security Engineering"
    assert doc.last_modified.strftime("%Y-%m-%d") == "2026-07-15"
    assert "---" not in doc.content
    assert "Body text here." in doc.content


def test_text_loader_parses_governance_frontmatter(tmp_path: Path):
    """Safety/freshness front-matter fields populate RawDocument's new governance fields."""
    path = tmp_path / "retention-policy-v2.md"
    path.write_text(
        "---\n"
        'title: "Tenant Alpha Retention Policy v2"\n'
        'tenant_id: "tenant_alpha"\n'
        "allowed_roles:\n"
        '  - "tenant_alpha_operator"\n'
        '  - "tenant_alpha_admin"\n'
        'classification: "internal"\n'
        'status: "active"\n'
        'updated_at: "2026-06-01T10:00:00+02:00"\n'
        'effective_from: "2026-06-01"\n'
        'document_version: "2.0"\n'
        "supersedes: retention-policy-v1.md\n"
        'trust_level: "authoritative"\n'
        'source_type: "controlled_internal"\n'
        "---\n\n"
        "# Retention\n\nBody text here.",
        encoding="utf-8",
    )

    doc = TextLoader().load(path)

    assert doc.tenant_id == "tenant_alpha"
    assert doc.allowed_roles == ["tenant_alpha_operator", "tenant_alpha_admin"]
    assert doc.classification == "internal"
    assert doc.status == "active"
    assert doc.effective_from.isoformat() == "2026-06-01"
    assert doc.document_version == "2.0"
    assert doc.supersedes_source == "retention-policy-v1.md"
    assert doc.trust_level == "authoritative"
    assert doc.doc_source_type == "controlled_internal"
    assert doc.source_type == "markdown"  # unaffected: the loader-type field, not doc_source_type
    assert doc.last_modified.isoformat() == "2026-06-01T10:00:00+02:00"


def test_text_loader_governance_fields_none_without_frontmatter(tmp_path: Path):
    """A document with no governance front matter leaves every new field unset."""
    path = tmp_path / "plain.md"
    path.write_text("# Title\n\nbody, no frontmatter here", encoding="utf-8")

    doc = TextLoader().load(path)

    assert doc.tenant_id is None
    assert doc.allowed_roles is None
    assert doc.trust_level is None
    assert doc.doc_source_type is None
    assert doc.supersedes_source is None


def test_text_loader_handles_markdown_without_frontmatter(tmp_path: Path):
    """Markdown without a front-matter block falls back to filename/raw content."""
    path = tmp_path / "plain.md"
    path.write_text("# Title\n\nbody, no frontmatter here", encoding="utf-8")

    doc = TextLoader().load(path)

    assert doc.title == "plain"  # falls back to filename stem
    assert doc.author is None
    assert doc.content == "# Title\n\nbody, no frontmatter here"


def test_text_loader_detects_markdown(tmp_path: Path):
    """A .md extension sets source_type to "markdown"."""
    path = tmp_path / "note.md"
    path.write_text("# Title\n\nbody", encoding="utf-8")

    doc = TextLoader().load(path)

    assert doc.source_type == "markdown"


def test_html_loader_extracts_title_and_text(tmp_path: Path):
    """HTML title/lang attributes and visible body text are extracted."""
    path = tmp_path / "page.html"
    path.write_text(
        "<html lang='en'><head><title>My Page</title></head><body><p>Hello there</p></body></html>",
        encoding="utf-8",
    )

    doc = HTMLLoader().load(path)

    assert doc.title == "My Page"
    assert "Hello there" in doc.content
    assert doc.language == "en"


def test_docx_loader_extracts_paragraphs(tmp_path: Path):
    """Paragraph text and the title core property are extracted from a .docx."""
    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.core_properties.title = "My Doc"
    document.save(str(path))

    doc = DocxLoader().load(path)

    assert "First paragraph" in doc.content
    assert "Second paragraph" in doc.content
    assert doc.title == "My Doc"


def test_pdf_loader_reads_blank_page_without_error(tmp_path: Path):
    """A PDF with no extractable text loads with empty content, not an error."""
    path = tmp_path / "doc.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with path.open("wb") as f:
        writer.write(f)

    doc = PDFLoader().load(path)

    assert doc.source_type == "pdf"
    # A blank page still gets its `<!--page:N-->` sentinel (needed so a later
    # page's number stays correct); no real text/structure was extracted.
    assert doc.content.strip() == "<!--page:1-->"


def test_registry_maps_extensions_to_loaders(tmp_path: Path):
    """get_loader returns the right Loader subclass per file extension."""
    assert isinstance(get_loader(tmp_path / "a.txt"), TextLoader)
    assert isinstance(get_loader(tmp_path / "a.md"), TextLoader)
    assert isinstance(get_loader(tmp_path / "a.html"), HTMLLoader)
    assert isinstance(get_loader(tmp_path / "a.docx"), DocxLoader)
    assert isinstance(get_loader(tmp_path / "a.pdf"), PDFLoader)


def test_registry_raises_for_unsupported_extension(tmp_path: Path):
    """get_loader raises ValueError for an unregistered extension."""
    with pytest.raises(ValueError, match="No loader registered"):
        get_loader(tmp_path / "a.xyz")


# --- Layout-aware DOCX extraction ------------------------------------------


def test_docx_loader_extracts_headings_as_section_markers(tmp_path: Path):
    """Title/Heading-styled paragraphs become #/## markers the chunker turns into section_path."""
    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("My Report", style="Title")
    document.add_paragraph("1. Scope", style="Heading 1")
    document.add_paragraph("This section explains the scope.")
    document.save(str(path))

    doc = DocxLoader().load(path)
    spans = StructuredMarkdownChunker().split(doc.content, source_type=doc.source_type)

    prose = [s for s in spans if (s.content_type or "prose") == "prose"]
    assert any(s.section_path == "My Report > 1. Scope" for s in prose)


def test_docx_loader_extracts_table_as_pipe_markdown(tmp_path: Path):
    """A docx table becomes a content_type='table' chunk after the chunker runs."""
    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("Intro", style="Heading 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "retries"
    table.cell(1, 1).text = "5"
    document.save(str(path))

    doc = DocxLoader().load(path)
    spans = StructuredMarkdownChunker().split(doc.content, source_type=doc.source_type)

    table_spans = [s for s in spans if s.content_type == "table"]
    assert len(table_spans) == 1
    assert "retries" in table_spans[0].text
    assert "5" in table_spans[0].text


def test_docx_loader_tracks_page_breaks(tmp_path: Path):
    """Manual page breaks increment ChunkSpan.page across the document."""
    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("Page one content", style="Heading 1")
    document.add_page_break()
    document.add_paragraph("Page two content", style="Heading 1")
    document.save(str(path))

    doc = DocxLoader().load(path)
    spans = StructuredMarkdownChunker().split(doc.content, source_type=doc.source_type)

    pages = {s.page for s in spans if "Page one" in s.text or "Page one content" in s.text}
    assert pages == {1}
    pages_two = {s.page for s in spans if "Page two content" in s.text}
    assert pages_two == {2}


def test_docx_loader_detects_monospace_paragraph_as_code(tmp_path: Path):
    """A paragraph whose runs are all a monospace font becomes a fenced code/config block."""
    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("Setup", style="Heading 1")
    code_paragraph = document.add_paragraph()
    run = code_paragraph.add_run("retry_seconds: 30")
    run.font.name = "Consolas"
    document.save(str(path))

    doc = DocxLoader().load(path)
    spans = StructuredMarkdownChunker().split(doc.content, source_type=doc.source_type)

    code_spans = [s for s in spans if s.content_type in ("code", "configuration")]
    assert len(code_spans) == 1
    assert "retry_seconds: 30" in code_spans[0].text


def test_docx_loader_resolves_image_to_existing_asset_and_folds_caption(tmp_path: Path):
    """An embedded image is referenced via source_anchor into a pre-existing assets/ folder."""
    assets_dir = tmp_path / "assets"
    assets_dir.mkdir()
    (assets_dir / "diagram-view-01.png").write_bytes(_tiny_png_bytes())

    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("Overview", style="Heading 1")
    document.add_picture(io.BytesIO(_tiny_png_bytes()))
    document.add_paragraph("Figure 1. Overview diagram.")
    document.save(str(path))

    doc = DocxLoader().load(path)
    spans = StructuredMarkdownChunker().split(doc.content, source_type=doc.source_type)

    image_spans = [s for s in spans if s.content_type == "image"]
    assert len(image_spans) == 1
    assert image_spans[0].source_anchor == "assets/diagram-view-01.png"
    assert "Figure 1. Overview diagram." in image_spans[0].text


def test_docx_loader_writes_fallback_asset_when_no_assets_folder_exists(tmp_path: Path):
    """With no sibling assets/ folder, the loader extracts and writes the image itself."""
    path = tmp_path / "standalone-report.docx"
    document = docx.Document()
    document.add_paragraph("Overview", style="Heading 1")
    document.add_picture(io.BytesIO(_tiny_png_bytes()))
    document.save(str(path))

    doc = DocxLoader().load(path)

    written = tmp_path / "assets" / "standalone-report-figure-01.png"
    assert written.is_file()
    assert "assets/standalone-report-figure-01.png" in doc.content


# --- resolve_image_asset: shared assets/ folder disambiguation -------------


def test_resolve_image_asset_matches_by_shared_filename_word(tmp_path: Path):
    """Two sibling documents sharing one assets/ folder each resolve to their own images.

    Regression test for a real bug found via integration testing: naive
    positional pairing across a shared assets/ folder handed every
    document the *first* document's images.
    """
    assets_dir = tmp_path / "assets"
    assets_dir.mkdir()
    (assets_dir / "capacity-view-01.png").touch()
    (assets_dir / "capacity-view-02.png").touch()
    (assets_dir / "incident-view-01.png").touch()

    doc_a = tmp_path / "ingestion-capacity-analysis.docx"
    doc_b = tmp_path / "incident-084-visual-review.pdf"

    assert resolve_image_asset(doc_a, 0, lambda: (b"", ".png")) == "assets/capacity-view-01.png"
    assert resolve_image_asset(doc_a, 1, lambda: (b"", ".png")) == "assets/capacity-view-02.png"
    assert resolve_image_asset(doc_b, 0, lambda: (b"", ".png")) == "assets/incident-view-01.png"


def test_resolve_image_asset_falls_back_to_all_candidates_when_no_word_matches(tmp_path: Path):
    """A folder with no filename-word overlap still resolves positionally, not empty."""
    assets_dir = tmp_path / "assets"
    assets_dir.mkdir()
    (assets_dir / "figure-01.png").touch()

    doc = tmp_path / "unrelated-name.docx"

    assert resolve_image_asset(doc, 0, lambda: (b"", ".png")) == "assets/figure-01.png"


def test_resolve_image_asset_extracts_bytes_when_no_assets_dir(tmp_path: Path):
    """With no assets/ directory at all, the factory's bytes are written to a new one."""
    doc = tmp_path / "report.pdf"

    result = resolve_image_asset(doc, 0, lambda: (b"raw-bytes", ".png"))

    assert result == "assets/report-figure-01.png"
    assert (tmp_path / "assets" / "report-figure-01.png").read_bytes() == b"raw-bytes"


# --- Layout-aware PDF extraction: helper-level (no real PDF needed) --------


def test_pdf_loader_group_paragraphs_detects_heading_by_font_size():
    """A short, large-font line becomes a heading; body-sized text stays prose."""
    loader = PDFLoader(heading_font_size_ratio=1.15, title_font_size_ratio=1.8)
    words = [
        {"text": "Title", "x0": 0, "x1": 10, "top": 0, "bottom": 10, "size": 24.0},
        {"text": "Words", "x0": 12, "x1": 20, "top": 0, "bottom": 10, "size": 24.0},
        {"text": "Regular", "x0": 0, "x1": 10, "top": 30, "bottom": 40, "size": 10.0},
        {"text": "paragraph", "x0": 12, "x1": 25, "top": 30, "bottom": 40, "size": 10.0},
        {"text": "text.", "x0": 27, "x1": 35, "top": 30, "bottom": 40, "size": 10.0},
    ]

    paragraphs = loader._group_paragraphs(words, body_size=10.0)

    kinds = [p["kind"] for p in paragraphs]
    assert "heading1" in kinds
    assert "prose" in kinds
    heading = next(p for p in paragraphs if p["kind"] == "heading1")
    assert heading["text"] == "Title Words"


def test_pdf_loader_group_paragraphs_detects_monospace_run_as_code():
    """Consecutive Courier-font lines are grouped into one code paragraph, not prose."""
    loader = PDFLoader()
    words = [
        {
            "text": "retry_seconds:",
            "x0": 0,
            "x1": 10,
            "top": 0,
            "bottom": 10,
            "size": 10.0,
            "fontname": "Courier",
        },
        {
            "text": "30",
            "x0": 12,
            "x1": 15,
            "top": 0,
            "bottom": 10,
            "size": 10.0,
            "fontname": "Courier",
        },
    ]

    paragraphs = loader._group_paragraphs(words, body_size=10.0)

    assert len(paragraphs) == 1
    assert paragraphs[0]["kind"] == "code"
    assert paragraphs[0]["text"] == "retry_seconds: 30"


def test_pdf_loader_render_table_produces_pipe_markdown():
    """_render_table turns a pdfplumber-Table-shaped object into a pipe table."""

    class FakeTable:
        def extract(self) -> list[list[str]]:
            return [["Name", "Value"], ["retries", "5"]]

    result = PDFLoader._render_table(FakeTable())

    assert result.splitlines()[0] == "| Name | Value |"
    assert "| retries | 5 |" in result
